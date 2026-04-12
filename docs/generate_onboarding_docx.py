"""Genera el documento Word del cuestionario de onboarding de Atendoo."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "ONBOARDING_INFORMACION_CLIENTES.docx")

def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elem)

def add_info_box(doc, text, bg_color="E8F4FD", border_color="2196F3", prefix="ℹ️"):
    """Añade un recuadro informativo (simula blockquote)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(f"{prefix} {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.font.italic = True
    set_cell_shading(cell, bg_color)
    doc.add_paragraph()  # spacer

def add_pre_filled_box(doc, text):
    """Añade un recuadro de datos pre-rellenados."""
    add_info_box(doc, f"[PRE-RELLENO — VERIFICA SI ES CORRECTO]\n{text}", bg_color="FFF3E0", border_color="FF9800", prefix="📋")

def add_warning_box(doc, text):
    """Añade un recuadro de advertencia."""
    add_info_box(doc, text, bg_color="FFEBEE", border_color="F44336", prefix="⚠️")

def add_recommendation_box(doc, text):
    """Añade un recuadro de recomendación."""
    add_info_box(doc, text, bg_color="E8F5E9", border_color="4CAF50", prefix="💡")

def add_question(doc, number, text, required=True):
    """Añade una pregunta con formato."""
    tag = "[obligatorio]" if required else "[opcional]"
    p = doc.add_paragraph()
    run = p.add_run(f"{number} ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run = p.add_run(f"  {tag}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) if required else RGBColor(0x66, 0x66, 0x66)
    run.bold = True

def add_hint(doc, text):
    """Añade texto de ayuda en cursiva."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def add_section_header(doc, text):
    """Añade un header de sección."""
    doc.add_page_break()
    heading = doc.add_heading(text, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

def add_subsection_header(doc, text):
    """Añade un sub-header."""
    heading = doc.add_heading(text, level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_option_block(doc, letter, title, description, pros, cons, extra_notes=None, recommended=False):
    """Añade un bloque de opción (A/B/C/D)."""
    label = f"Opción {letter} — {title}"
    if recommended:
        label += " (RECOMENDADA)"
    h = doc.add_heading(label, level=3)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E) if recommended else RGBColor(0x33, 0x33, 0x33)

    p = doc.add_paragraph()
    run = p.add_run("Cómo funciona: ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(description)
    run.font.size = Pt(10)

    # Pros
    p = doc.add_paragraph()
    run = p.add_run("✅ Ventajas:")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    for pro in pros:
        p = doc.add_paragraph(pro, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)

    # Cons
    p = doc.add_paragraph()
    run = p.add_run("❌ Inconvenientes:")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    for con in cons:
        p = doc.add_paragraph(con, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)

    if extra_notes:
        for note in extra_notes:
            p = doc.add_paragraph()
            run = p.add_run(note)
            run.font.size = Pt(10)
            run.font.italic = True

def add_service_block(doc, num, name, data):
    """Añade un bloque de servicio con los campos pre-rellenados."""
    doc.add_heading(f"Servicio {num}: {name}", level=3)
    for field, value in data:
        p = doc.add_paragraph()
        run = p.add_run(f"{field}: ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(str(value))
        run.font.size = Pt(10)
        if value and ("Sin datos" in str(value) or "CONFIRMAR" in str(value)):
            run.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)

def build_document():
    doc = Document()

    # ── Estilos globales ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # ═══════════════════════════════════════
    # PORTADA
    # ═══════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ATENDOO")
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Cuestionario de Onboarding")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Toda la información que tu bot de WhatsApp necesita\npara funcionar correctamente en tu negocio")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Tiempo estimado: 15-25 minutos")
    run.font.size = Pt(11)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Los campos marcados con ")
    run.font.size = Pt(10)
    run = p.add_run("[obligatorio]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.bold = True
    run = p.add_run(" son imprescindibles. Los marcados con ")
    run.font.size = Pt(10)
    run = p.add_run("[opcional]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.bold = True
    run = p.add_run(" puedes dejarlos en blanco si no aplican.")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nLos recuadros naranjas contienen datos que ya tenemos pre-rellenados.\nSolo necesitas verificar si son correctos o corregirlos.")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0xFF, 0x98, 0x00)

    # ═══════════════════════════════════════
    # §1. IDENTIDAD Y CONTACTO
    # ═══════════════════════════════════════
    add_section_header(doc, "§1. Identidad y contacto del negocio")

    add_question(doc, "1.1", "Nombre comercial del negocio")
    add_hint(doc, "(ej: Clínica Physiofitness, Peluquería María, Centro Deportivo FitZone)")
    add_pre_filled_box(doc, "Clínica Physiofitness")

    add_question(doc, "1.2", "Razón social / NIF", required=False)
    add_hint(doc, "(ej: Physiofitness S.L., NIF: B12345678. Si eres autónomo, déjalo en blanco)")
    add_pre_filled_box(doc, "(Sin datos — completar si aplica)")

    add_question(doc, "1.3", "Dirección física exacta")
    add_hint(doc, "(ej: Calle Venezuela, 8, 28220 Majadahonda, Madrid)")
    add_pre_filled_box(doc, "Calle Venezuela, 8, 28220 Majadahonda, Madrid")

    add_question(doc, "1.4", "URL de Google Maps", required=False)
    add_hint(doc, "(Pega aquí el enlace directo a tu ubicación en Google Maps)")
    add_pre_filled_box(doc, "(Sin datos — pegar enlace de Google Maps)")

    add_question(doc, "1.5", "Teléfono principal del negocio")
    add_hint(doc, "(ej: 618 974 833)")
    add_pre_filled_box(doc, "618 974 833")

    add_question(doc, "1.6", "¿Tienes un número de teléfono fijo en el negocio?", required=False)
    add_hint(doc, '(ej: 91 123 45 67. Si no tienes fijo, escribe "No")')
    add_pre_filled_box(doc, "(Sin datos — ¿tienes fijo?)")

    add_question(doc, "1.7", "Email principal del negocio")
    add_hint(doc, "(A este email llegarán las notificaciones del bot)")
    add_pre_filled_box(doc, "david.fisiofit@gmail.com")

    add_question(doc, "1.8", "Email de notificaciones del bot", required=False)
    add_hint(doc, "(Si quieres que las alertas del bot lleguen a un email diferente al principal)")
    add_pre_filled_box(doc, "(Mismo que el principal — dejar en blanco si no cambia)")

    add_question(doc, "1.9", "Web del negocio", required=False)
    add_pre_filled_box(doc, "https://www.clinicaphysiofitness.com/")

    add_question(doc, "1.10", "Redes sociales", required=False)
    add_hint(doc, "(Indica los perfiles activos: Instagram, Facebook, etc.)")
    add_pre_filled_box(doc, "(Sin datos — completar las redes activas)")

    # ═══════════════════════════════════════
    # §2. WHATSAPP
    # ═══════════════════════════════════════
    add_section_header(doc, "§2. Decisión de infraestructura WhatsApp")

    add_warning_box(doc, (
        "BLOQUE CRÍTICO — Lee con atención antes de responder\n\n"
        "Para que el bot de Atendoo funcione por WhatsApp, necesita un número de teléfono dedicado, "
        "por requerimientos obligatorios de WhatsApp Business.\n\n"
        "¿Qué significa esto? Cuando un número de teléfono se registra en WhatsApp Business, "
        "deja de funcionar en la app normal de WhatsApp (tanto la personal como la de WhatsApp Business App). Es decir:\n"
        "• Ya no puedes abrir WhatsApp en el móvil con ese número.\n"
        "• Pierdes acceso al historial de chats que tenías.\n"
        "• Sales de todos los grupos de WhatsApp.\n"
        "• Ya no puedes enviar ni recibir mensajes desde la app.\n"
        "• Las llamadas telefónicas normales (voz) SÍ siguen funcionando.\n"
        "• Los datos de WhatsApp (chats, grupos, historial) se pierden de forma permanente. "
        "Incluso si en el futuro se des-registra el número de la API y se vuelve a usar WhatsApp normal, "
        "se empezaría de cero — no se recupera nada de lo anterior.\n\n"
        "Por eso, elegir qué número usa el bot es una decisión importante."
    ))

    add_question(doc, "2.1", "¿El número principal del negocio (618 974 833) se usa actualmente para WhatsApp?")
    add_hint(doc, "(Responde Sí o No)")
    add_pre_filled_box(doc, "(Suponemos que sí — confirmar)")

    add_question(doc, "2.2", "¿Usas ese WhatsApp para comunicarte con pacientes, proveedores o compañeros?")
    add_hint(doc, '(Responde Sí o No. Si Sí, indica brevemente para qué)')
    add_pre_filled_box(doc, "(Sin datos — necesitamos saberlo para recomendarte la mejor opción)")

    add_question(doc, "2.3", "¿Tienes algún número de teléfono (móvil o fijo) que NO estés usando con WhatsApp?")
    add_hint(doc, (
        "Piensa si tienes: un número fijo de la clínica, una SIM antigua en un cajón, "
        "un segundo número móvil, un número de un móvil viejo... "
        "Si tienes alguno, lo podemos usar para el bot y te ahorras el coste de contratar uno nuevo."
    ))
    add_pre_filled_box(doc, '(Sin datos — escribe el número o "No tengo ninguno disponible")')

    add_question(doc, "2.4", "Elige la opción que prefieras para el número del bot")
    p = doc.add_paragraph("Lee las 4 opciones con calma. Cada una tiene sus ventajas e inconvenientes.")
    p.runs[0].font.size = Pt(10)

    # Opción A
    add_option_block(doc, "A", "Migrar tu número actual (618 974 833) al bot",
        "Tu número actual se registra en la API de WhatsApp Business. El bot responde a los pacientes desde ese mismo número. Sigues recibiendo llamadas normales sin problema.",
        [
            "Los pacientes ya conocen tu número — no hay que comunicar ningún cambio.",
            "Coste adicional: 0 €.",
            "Máxima continuidad: el bot \"hereda\" la identidad de contacto de la clínica.",
        ],
        [
            "Pierdes el acceso a WhatsApp en ese número. Nada de chats, grupos, ni mensajes — ni los antiguos ni nuevos.",
            "Si usas ese WhatsApp para hablar con pacientes, proveedores, amigos o familia, todo eso se pierde.",
            "El historial de conversaciones no se transfiere al bot — se pierde al eliminar la cuenta.",
            "Si en algún momento quisieras volver a usar WhatsApp normal en ese número, los chats, grupos e historial anteriores NO se recuperan.",
        ],
        [
            "💾 Si eliges esta opción, haz una copia de seguridad ANTES: WhatsApp → Ajustes → Chats → Copia de seguridad (guarda en Google Drive / iCloud). También puedes exportar chats individuales: chat → ⋮ → Más → Exportar chat.",
            "⚠️ Solo elige esta opción si no usas WhatsApp en el 618 974 833 para nada personal, o si no te importa perder ese acceso.",
        ],
    )

    # Opción B
    add_option_block(doc, "B", "Contratar un número nuevo para el bot",
        "Se compra un número de teléfono nuevo (virtual o SIM prepago) exclusivo del bot. Tu WhatsApp actual NO se toca — sigue todo igual.",
        [
            "No pierdes absolutamente nada. Tu WhatsApp del 618 974 833 sigue funcionando tal cual.",
            "Separación limpia: el bot tiene su número, tú tienes el tuyo.",
            "Si el bot necesita derivar, te avisa al 618 974 833 y tú respondes desde tu WhatsApp normal.",
            "Fácil de comunicar: se pone el número del bot en la web, Instagram, Google Business.",
        ],
        [
            "Hay un coste mensual: ~5-15 €/mes (SIM prepago o número virtual VoIP).",
            "Los pacientes ven un número nuevo que no conocen (se comunica una vez y listo).",
        ],
        [
            "💡 Para un negocio que factura a 55 €/sesión, 5-15 €/mes es menos que el precio de una sesión.",
            "📱 Nosotros podemos encargarnos: contratamos el número, lo configuramos y lo dejamos listo.",
        ],
        recommended=True,
    )

    # Opción C
    add_option_block(doc, "C", "Usar un número que ya tengas y que no esté en WhatsApp",
        "Si tienes algún número (móvil o fijo) que NO esté registrado en WhatsApp, lo usamos para el bot. No compras nada nuevo.",
        [
            "Coste adicional: 0 € — usamos un número que ya tienes.",
            "Tu WhatsApp del 618 974 833 sigue intacto.",
            "Separación limpia entre bot y uso personal.",
        ],
        [
            "El número debe poder recibir un SMS o llamada de verificación de Meta (una sola vez).",
            "Si es un número fijo, los pacientes verán un fijo en WhatsApp (cada vez más común en negocios).",
            "Si es una SIM antigua, debe estar activa y no caducada.",
        ],
        [
            "📋 Ejemplos: fijo de la clínica, SIM de prepago antigua, móvil viejo con su línea, número de empresa secundario.",
            "⚠️ El número NO debe estar registrado en ninguna app de WhatsApp actualmente.",
        ],
    )

    # Opción D
    add_option_block(doc, "D", "Usar el número fijo de la clínica",
        "Si la clínica tiene un fijo (tipo 91X XXX XXX), lo registramos en la API de WhatsApp Business. La verificación se hace por llamada de voz.",
        [
            "Coste adicional: 0 € — ya estás pagando la línea fija.",
            "Tu WhatsApp del 618 974 833 sigue intacto.",
            "El fijo puede ser reconocible como \"el teléfono de la clínica\".",
            "Muchos negocios con bot usan el fijo para WhatsApp y el móvil para llamadas personales.",
        ],
        [
            "Los pacientes ven un número fijo en WhatsApp (cada vez más común).",
            "El fijo NO debe tener WhatsApp previamente registrado.",
            "Si el fijo tiene centralita o filtros de llamadas, la verificación de Meta podría no entrar.",
            "El fijo queda reservado para el bot (no podrás usarlo para WhatsApp personal).",
        ],
        ["⚠️ Necesitas poder contestar el fijo cuando Meta llame para verificar (una sola llamada)."],
    )

    # Tabla resumen
    doc.add_heading("Resumen de opciones", level=3)
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["", "Opción A", "Opción B", "Opción C", "Opción D"]
    rows_data = [
        ["Número del bot", "Tu 618 974 833", "Número nuevo", "Un nº tuyo sin WA", "Tu fijo"],
        ["Conservas tu WhatsApp", "❌ No", "✅ Sí", "✅ Sí", "✅ Sí"],
        ["Pacientes reconocen nº", "✅ Sí", "❌ No (es nuevo)", "Depende", "Si conocen el fijo"],
        ["Coste extra mensual", "0 €", "5-15 €", "0 €", "0 €"],
        ["Riesgo", "🔴 Alto", "🟢 Bajo", "🟢 Bajo", "🟡 Medio"],
    ]

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    for row_idx, row_data in enumerate(rows_data, 1):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()
    add_recommendation_box(doc, (
        "Recomendamos B si usas activamente tu WhatsApp actual, o C/D si tienes un número disponible "
        "(fijo o móvil) que no estés usando con WhatsApp — así te ahorras el coste mensual."
    ))

    add_question(doc, "2.5", "Si elegiste B: ¿prefieres que contratemos nosotros el número o lo gestionas tú?", required=False)
    add_question(doc, "2.6", "Si elegiste C: ¿cuál es el número que quieres usar?", required=False)
    add_question(doc, "2.7", "Si elegiste D: ¿cuál es el número fijo de la clínica?", required=False)
    add_question(doc, "2.8", "¿Tu teléfono móvil personal tiene dual SIM o eSIM?", required=False)

    # ═══════════════════════════════════════
    # §3. HORARIO
    # ═══════════════════════════════════════
    add_section_header(doc, "§3. Horario de operaciones")

    add_question(doc, "3.1", "Horario de atención por día")
    add_hint(doc, "(Indica el horario para cada día. Si tienes mañana y tarde separados, indícalo)")

    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light List Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = "Día"
    table.rows[0].cells[1].text = "Horario"
    days = [
        ("Lunes", "07:00 - 21:00"), ("Martes", "07:00 - 21:00"),
        ("Miércoles", "07:00 - 21:00"), ("Jueves", "07:00 - 21:00"),
        ("Viernes", "07:00 - 21:00"), ("Sábado", "10:00 - 14:00"),
        ("Domingo", "Cerrado"),
    ]
    for i, (day, hours) in enumerate(days, 1):
        table.rows[i].cells[0].text = day
        table.rows[i].cells[1].text = hours

    doc.add_paragraph()
    add_info_box(doc, "Pregunta adicional: ¿El horario L-V es continuo (7 a 21 sin pausa) o hay un descanso a mediodía? Esto es importante para que el bot ofrezca huecos correctamente.")

    add_question(doc, "3.2", "Festivos y vacaciones", required=False)
    add_hint(doc, '(ej: "Cerramos del 1 al 15 de agosto", "Los festivos de la Comunidad de Madrid no abrimos")')
    add_pre_filled_box(doc, "(Sin datos — completar)")

    add_question(doc, "3.3", "¿El bot puede ofrecer citas en cualquier franja del horario, o hay restricciones?")
    add_hint(doc, '(ej: "Solo ofrezco fisioterapia por las mañanas", "Todo igual durante todo el horario")')
    add_pre_filled_box(doc, "(Sin datos — ¿hay restricciones horarias por tipo de servicio?)")

    add_question(doc, "3.4", "¿Qué debe hacer el bot cuando alguien quiere cita fuera de horario?")
    p = doc.add_paragraph("□ Rechazar la solicitud y ofrecer el próximo hueco disponible")
    p = doc.add_paragraph("□ Añadir a una lista de espera")
    p = doc.add_paragraph("□ Derivar al profesional para que decida")
    p = doc.add_paragraph("□ El bot funciona 24/7 pero solo ofrece huecos dentro del horario")
    add_recommendation_box(doc, "Recomendamos: El bot funciona 24/7 pero solo ofrece huecos dentro del horario.")

    # ═══════════════════════════════════════
    # §4. CATÁLOGO DE SERVICIOS
    # ═══════════════════════════════════════
    add_section_header(doc, "§4. Catálogo de servicios")
    add_warning_box(doc, "BLOQUE CLAVE — Esta sección es fundamental para que el bot ofrezca tus servicios correctamente. Rellena CADA servicio por separado.")

    services = [
        ("Fisioterapia", [
            ("Categoría", "Fisioterapia"), ("Duración", "50 minutos"), ("Precio sesión suelta", "55 €"),
            ("Bonos", "Bono 5: 265 € (53 €/sesión) | Bono 10: 500 € (50 €/sesión)"),
            ("Suscripción mensual", "No"), ("Profesional", "Sin datos — ¿quién lo realiza?"),
            ("Requiere valoración previa", "Sin datos — CONFIRMAR"), ("Recurso/máquina", "Sin datos"),
            ("Individual o grupal", "Individual"),
        ]),
        ("Entrenamiento Personal", [
            ("Categoría", "Entrenamiento"), ("Duración", "60 minutos"), ("Precio sesión suelta", "50 €"),
            ("Bonos", "Bono 5: 240 € (48 €/sesión) | Bono 10: 450 € (45 €/sesión)"),
            ("Suscripción mensual", "No"), ("Profesional", "Sin datos — ¿quién lo realiza?"),
            ("Requiere valoración previa", "Sin datos"), ("Recurso/máquina", "No"),
            ("Individual o grupal", "Individual"),
        ]),
        ("Entrenamiento Grupos Reducidos", [
            ("Categoría", "Entrenamiento"), ("Duración", "60 minutos"), ("Precio sesión suelta", "55 € / persona"),
            ("Bonos", "No (funciona con suscripciones)"),
            ("Suscripciones", "1 clase/semana: 140 €/mes | 2 clases: 260 €/mes | 3 clases: 360 €/mes"),
            ("Profesional", "Sin datos"), ("Requiere valoración previa", "Sin datos"),
            ("Recurso/máquina", "No"), ("Individual o grupal", "Grupal — 2 a 4 personas"),
        ]),
        ("Fisio-estética 1 zona (30 min)", [
            ("Categoría", "Estética"), ("Duración", "30 minutos"), ("Precio sesión suelta", "55 €"),
            ("Bonos", "Bono 5: 250 € | Bono 10: 500 €"),
            ("Zonas disponibles", "Facial, Piernas, Abdomen, Glúteos"),
            ("Profesional", "Sin datos"), ("Recurso/máquina", "Sin datos — ¿INDIBA?"),
            ("Individual o grupal", "Individual"),
        ]),
        ("Fisio-estética 1 zona (45 min)", [
            ("Categoría", "Estética"), ("Duración", "45 minutos"), ("Precio sesión suelta", "70 €"),
            ("Bonos", "Bono 5: 340 € | Bono 10: 650 €"),
            ("Zonas disponibles", "Facial, Piernas, Abdomen, Glúteos"),
            ("Individual o grupal", "Individual"),
        ]),
        ("Fisio-estética 2 zonas (60 min)", [
            ("Categoría", "Estética"), ("Duración", "60 minutos"), ("Precio sesión suelta", "90 €"),
            ("Bonos", "Bono 5: 425 € | Bono 10: 800 €"),
            ("Individual o grupal", "Individual"),
        ]),
        ("Fisio-estética 3 zonas (60 min)", [
            ("Categoría", "Estética"), ("Duración", "60 minutos"), ("Precio sesión suelta", "99 €"),
            ("Bonos", "Bono 5: 475 € | Bono 10: 900 €"),
            ("Individual o grupal", "Individual"),
        ]),
        ("Presoterapia", [
            ("Categoría", "Complementario"), ("Duración", "30 minutos"), ("Precio sesión suelta", "20 €"),
            ("Bonos", "Sin datos — ¿ofreces bonos?"),
            ("Recurso/máquina", "Sí — máquina de presoterapia"),
            ("Individual o grupal", "Individual"),
        ]),
        ("Readaptación Neuromuscular", [
            ("Categoría", "Complementario"), ("Duración", "Sin datos — ¿cuántos minutos?"),
            ("Precio sesión suelta", "60 €"), ("Bonos", "Sin datos — ¿ofreces bonos?"),
            ("Recurso/máquina", "Sí — máquina isocinética"),
            ("Individual o grupal", "Individual"),
        ]),
    ]

    for i, (name, data) in enumerate(services, 1):
        add_service_block(doc, i, name, data)

    p = doc.add_paragraph()
    run = p.add_run("\n¿Tienes más servicios? Si es así, descríbelos con el mismo formato.")
    run.bold = True
    run.font.size = Pt(10)

    # ═══════════════════════════════════════
    # §5. RESTRICCIONES MÉDICAS
    # ═══════════════════════════════════════
    add_section_header(doc, "§5. Restricciones médicas y límites del bot")
    add_warning_box(doc, "BLOQUE CRÍTICO — Esta sección define qué puede y qué NO puede decir tu bot. Tómate tu tiempo.")

    add_question(doc, "5.1", "¿Qué patologías o problemas tratas habitualmente?")
    add_hint(doc, "(ej: cervicalgias, lumbalgias, ciáticas, contracturas, tendinitis, post-operatorio, fascitis plantar...)")
    add_pre_filled_box(doc, "(Sin datos — necesitamos que listes tus patologías habituales)")

    add_question(doc, "5.2", "¿Qué patologías NO tratas o requieren derivación inmediata?")
    add_hint(doc, "(ej: sospecha de fractura, traumatismos <24h, dolor nocturno irradiado, fiebre + dolor articular, pérdida de fuerza, déficit neurológico...)")
    add_pre_filled_box(doc, "(Sin datos — FUNDAMENTAL definir las líneas rojas)")

    add_question(doc, "5.3", "Contraindicaciones para tratamientos estéticos")
    add_hint(doc, "(ej: embarazo, marcapasos, prótesis metálicas, problemas circulatorios, cáncer activo, heridas abiertas...)")
    add_pre_filled_box(doc, "(Sin datos — listar contraindicaciones de INDIBA/presoterapia)")

    add_question(doc, "5.4", "¿El bot atiende menores de edad?")
    add_hint(doc, '(ej: "Sí, cualquier edad" / "Solo mayores de 16" / "Menores con autorización de los padres")')

    add_question(doc, "5.5", "¿Tratas pacientes de edad avanzada? ¿Hay alguna limitación?", required=False)

    add_question(doc, "5.6", "¿El bot puede dar consejos clínicos o solo gestionar la agenda?")
    p = doc.add_paragraph("(A) Solo agenda: el bot NO opina sobre dolencias.")
    p = doc.add_paragraph("(B) Mini-triaje empático (RECOMENDADA): el bot hace 1-2 preguntas de triaje y da respuesta empática, nunca da diagnóstico.")
    p = doc.add_paragraph("(C) Información detallada: el bot da explicaciones amplias. ⚠️ Mayor riesgo legal.")
    add_recommendation_box(doc, "Recomendamos B — minimiza derivaciones innecesarias y mejora la conversión a cita, sin riesgo legal.")

    add_question(doc, "5.7", "¿El bot puede mencionar precios concretos?")
    add_hint(doc, '(ej: "Sí" / "No, que contacte con la clínica" / "Solo el precio de sesión suelta")')

    # ═══════════════════════════════════════
    # §6. TONO Y PERSONALIDAD
    # ═══════════════════════════════════════
    add_section_header(doc, "§6. Tono y personalidad del bot")

    add_question(doc, "6.1", "¿El bot tutea o trata de usted?")
    add_question(doc, "6.2", "Nivel de cercanía")
    p = doc.add_paragraph("• Clínico-formal  • Cercano-empático (RECOMENDADA)  • Desenfadado")
    add_question(doc, "6.3", "¿El bot puede usar emojis?")
    add_question(doc, "6.4", "Idioma principal del bot")
    add_pre_filled_box(doc, "Solo español")
    add_question(doc, "6.5", "¿Hay alguna palabra o frase que el bot NO debe usar nunca?", required=False)
    add_question(doc, "6.6", "¿Quieres que el bot firme sus mensajes?", required=False)

    # ═══════════════════════════════════════
    # §7. POLÍTICA DE CITAS
    # ═══════════════════════════════════════
    add_section_header(doc, "§7. Política de citas")

    add_question(doc, "7.1", "¿Se puede reservar cita directamente la primera vez o se requiere valoración previa?")
    add_question(doc, "7.2", "¿Con cuánta antelación mínima se puede cancelar sin coste?")
    add_pre_filled_box(doc, "24 horas")
    add_question(doc, "7.3", "¿Qué pasa si un paciente no se presenta (no-show)?")
    add_question(doc, "7.4", "¿Tienes un enlace de reservas online externo?", required=False)
    add_question(doc, "7.5", "¿El bot puede ofrecer el enlace de reservas como alternativa?", required=False)

    # ═══════════════════════════════════════
    # §8. DATOS DEL PROFESIONAL
    # ═══════════════════════════════════════
    add_section_header(doc, "§8. Datos del profesional para derivaciones")

    add_question(doc, "8.1", "Nombre completo del profesional principal")
    add_hint(doc, '(El bot usará este nombre: "lo confirma [nombre] en tu cita")')
    add_pre_filled_box(doc, "David (falta apellido — completar)")

    add_question(doc, "8.2", "WhatsApp personal del profesional")
    add_pre_filled_box(doc, "618 974 833")

    add_question(doc, "8.3", "Email para notificaciones de derivación")
    add_pre_filled_box(doc, "david.fisiofit@gmail.com")

    add_question(doc, "8.4", "Horario de disponibilidad para responder derivaciones", required=False)
    add_question(doc, "8.5", "¿Hay más profesionales en el equipo?", required=False)

    # ═══════════════════════════════════════
    # §9. RGPD
    # ═══════════════════════════════════════
    add_section_header(doc, "§9. RGPD y consentimiento")

    add_question(doc, "9.1", "Texto del aviso de privacidad para el primer mensaje")
    add_pre_filled_box(doc, "Al continuar, aceptas que procesemos tus datos para gestionar tu cita.")

    add_question(doc, "9.2", "URL de la política de privacidad completa")
    add_pre_filled_box(doc, "(Sin datos — ¿tienes política de privacidad publicada?)")

    add_question(doc, "9.3", "¿Cómo se pide el consentimiento?")
    p = doc.add_paragraph("(A) Consentimiento explícito: el paciente responde \"Acepto\".")
    p = doc.add_paragraph("(B) Consentimiento implícito (RECOMENDADA): al continuar la conversación se entiende como aceptación.")
    add_recommendation_box(doc, "Recomendamos B — menos fricción, estándar en la industria.")

    # ═══════════════════════════════════════
    # §10. EDGE CASES
    # ═══════════════════════════════════════
    add_section_header(doc, "§10. Casos límite (edge-cases)")
    add_info_box(doc, 'Si no tienes claro alguna respuesta, escribe "No sé — decidir juntos".')

    add_question(doc, "10.1", "¿Qué hace el bot si un paciente describe síntomas graves?")
    add_recommendation_box(doc, 'Recomendamos: "Por lo que me cuentas, prefiero que [nombre] te valore cuanto antes. Te paso con él/ella ahora mismo." y derivar.')

    add_question(doc, "10.2", "¿Y si un paciente quiere reservar para un menor de edad?")
    add_question(doc, "10.3", "¿Y si un paciente pregunta por servicios que no ofrecéis?")
    add_question(doc, "10.4", "¿Trabajáis con mutuas o seguros médicos?")
    add_question(doc, "10.5", "¿Y si un paciente quiere hablar con una persona concreta?")
    add_question(doc, "10.6", "¿Y si un paciente pide cambiar de profesional?", required=False)

    add_question(doc, "10.7", "Mensaje cuando el bot no entiende")
    add_recommendation_box(doc, 'Recomendamos: "Disculpa, no he entendido bien 😅 ¿Puedes decírmelo de otra forma? Si prefieres, te pongo directamente con [nombre]."')

    add_question(doc, "10.8", "Preguntas frecuentes que debería manejar el bot", required=False)
    add_hint(doc, "Lista las preguntas que tus pacientes hacen repetidamente:")
    p = doc.add_paragraph("1. ¿Necesito cita previa? →")
    p = doc.add_paragraph("2. ¿Cuánto dura una sesión? →")
    p = doc.add_paragraph("3. ¿Hacéis primera consulta gratuita? →")
    p = doc.add_paragraph("4. ¿Dónde puedo aparcar? →")
    p = doc.add_paragraph("5. ¿Qué debo traer a la primera cita? →")
    p = doc.add_paragraph("(Añade las que falten)")

    # ═══════════════════════════════════════
    # GUARDAR
    # ═══════════════════════════════════════
    doc.save(OUTPUT)
    print(f"Documento generado: {OUTPUT}")

if __name__ == "__main__":
    build_document()
